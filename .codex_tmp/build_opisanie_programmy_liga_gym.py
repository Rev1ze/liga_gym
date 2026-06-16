from pathlib import Path

from docx import Document
from docx.shared import Pt


OUT = Path("docs/opisanie_programmy_liga_gym.docx")


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, bold=bold)
    return p


def add_table(doc, caption, rows):
    add_paragraph(doc, caption)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Процедура / Модуль"
    table.rows[0].cells[1].text = "Описание процедуры"

    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, size=11, bold=True)

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=11)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    add_paragraph(doc, "2 Экспериментальный раздел", bold=True)
    add_paragraph(doc, "2.1 Описание программы", bold=True)
    add_paragraph(
        doc,
        "Управление состоянием реализовано через flutter_riverpod и "
        "riverpod_annotation. Для каждого крупного функционального блока выделены "
        "провайдеры, репозитории, источники данных и контроллеры экранов, что "
        "отделяет интерфейс от бизнес-логики и хранения данных.",
    )
    add_paragraph(
        doc,
        "В рамках разработки программного продукта реализовано единое мобильное "
        "приложение Liga Gym, предназначенное для авторизации пользователя, учета "
        "тренировок, маршрутов, питания, шаговой активности, целей, аналитики "
        "прогресса, общения с друзьями, работы с тренером и получения "
        "AI-рекомендаций по фитнесу.",
    )
    add_paragraph(
        doc,
        "Клиентская часть реализована на языке Dart с использованием фреймворка "
        "Flutter и библиотеки Material 3. Интерфейс приложения включает splash-экран "
        "с Lottie-анимацией, экраны входа и регистрации, анкету профиля, главный "
        "экран, тренировочные экраны и формы работы с дневником питания. Для "
        "локализации используются ресурсы на русском и английском языках.",
    )
    add_paragraph(
        doc,
        "Для удаленного хранения данных и аутентификации используются Firebase "
        "Authentication, Google Sign-In и Cloud Firestore. Локальное хранение "
        "тренировок, записей питания и шагов организовано на базе SQLite через "
        "библиотеку sqflite. Маршрут тренировки формируется на основе данных "
        "Geolocator, а поиск продукта по штрихкоду выполняется по встроенному "
        "каталогу продуктов.",
    )
    add_paragraph(
        doc,
        "Архитектура приложения построена по модульному принципу с выделением "
        "функциональных областей auth, dashboard, workout, nutrition, steps, "
        "ai_coach, social, coach, notifications, water_tracker и workout_completion. "
        "Внутри модулей код разделен на уровни presentation, domain и data.",
    )
    add_paragraph(
        doc,
        "Модульная схема проекта представлена в приложении А.",
    )
    add_paragraph(
        doc,
        "В таблице 2.1.1 описывается каждый модуль, перечисляются процедуры, "
        "входящие в модуль, и действия, выполняемые в каждой процедуре модуля.",
    )

    rows = [
        ("app.dart", "Точка сборки мобильного приложения, подключает локализацию, тему, маршрутизацию и основные провайдеры."),
        ("LigaGymApp", "Класс корневого виджета приложения, формирующий MaterialApp и общую структуру интерфейса."),
        ("LigaGymApp.build()", "Формирует корневой интерфейс приложения, подключает локализации, тему и генератор маршрутов."),
        ("core/navigation/app_router.dart", "Формирует экраны по именам маршрутов и передает аргументы навигации."),
        ("onGenerateRoute()", "Создает и возвращает нужный экран приложения по имени маршрута."),
        ("core/navigation/app_routes.dart", "Содержит константы именованных маршрутов мобильного приложения."),
        ("core/localization", "Содержит ресурсы локализации и расширения для вывода интерфейса на русском и английском языках."),
        ("core/supabase/supabase_bootstrap.dart", "Инициализирует Supabase и определяет bucket coach-media для тренерских медиафайлов."),
        ("core/notifications/app_notification_service.dart", "Настраивает локальные уведомления, включая напоминания перед запланированными тренировками."),
        ("core/offline", "Содержит общие интерфейсы для синхронизации локальных записей с удаленным хранилищем."),
        ("features/auth", "Модуль авторизации, регистрации, входа через Google, заполнения и редактирования профиля."),
        ("AuthActionController", "Контроллер состояния и действий модуля авторизации."),
        ("AuthActionController.loginWithEmail()", "Выполняет вход пользователя по email и паролю через Firebase Authentication."),
        ("AuthActionController.signInWithGoogle()", "Запускает сценарий входа через учетную запись Google."),
        ("AuthActionController.registerUser()", "Создает новую учетную запись пользователя и возвращает результат авторизации."),
        ("AuthActionController.saveUserProfile()", "Сохраняет заполненный профиль пользователя после регистрации."),
        ("AuthActionController.signOut()", "Завершает пользовательскую сессию и очищает auth-состояние."),
        ("SplashController.checkUserAuthState()", "Проверяет состояние авторизации и определяет целевой экран приложения."),
        ("auth_providers.dart", "Создает провайдеры FirebaseAuth, GoogleSignIn, FirebaseFirestore, репозитория и сценариев авторизации."),
        ("authStateChanges()", "Подписывается на поток изменений состояния авторизации Firebase."),
        ("currentAuthUser()", "Возвращает текущего авторизованного пользователя."),
        ("LoginScreen", "Экран входа пользователя по email и паролю или через Google."),
        ("LoginScreen._handleEmailLogin()", "Проверяет введенные данные и вызывает вход по email."),
        ("LoginScreen._handleGoogleSignIn()", "Запускает авторизацию через Google."),
        ("RegisterScreen", "Экран регистрации нового пользователя."),
        ("RegisterScreen._handleRegister()", "Передает данные регистрации в контроллер авторизации."),
        ("ProfileSetupScreen", "Экран первичного заполнения пользовательского профиля."),
        ("ProfileSetupScreen._pickBirthDate()", "Открывает выбор даты рождения пользователя."),
        ("ProfileSetupScreen._handleSaveProfile()", "Сохраняет имя, пол, дату рождения, цели и параметры профиля."),
        ("ProfileScreen", "Экран просмотра и редактирования профиля, целей, веса и тренерских запросов."),
        ("ProfileScreen._handleSave()", "Обновляет данные профиля пользователя в Firestore."),
        ("ProfileScreen._handleSignOut()", "Завершает текущую пользовательскую сессию."),
        ("ProfileScreen._accept()", "Принимает входящий запрос тренера."),
        ("ProfileScreen._decline()", "Отклоняет входящий запрос тренера."),
        ("FirestoreProfileRemoteDataSource", "Источник данных профиля, работающий с коллекцией users в Cloud Firestore."),
        ("FirestoreProfileRemoteDataSource.saveUserProfile()", "Сохраняет профиль пользователя, код друга, цели и историю веса."),
        ("AuthRepositoryImpl", "Репозиторий авторизации, объединяющий FirebaseAuth, GoogleSignIn и Firestore."),
        ("AuthRepositoryImpl.checkUserAuthState()", "Проверяет наличие пользователя и заполненность профиля."),
        ("AuthRepositoryImpl.loginWithEmail()", "Выполняет вход по email и паролю."),
        ("AuthRepositoryImpl.signInWithGoogle()", "Выполняет вход через Google."),
        ("AuthRepositoryImpl.registerUser()", "Создает учетную запись пользователя."),
        ("AuthRepositoryImpl.saveUserProfile()", "Передает данные профиля в удаленный источник данных."),
        ("features/dashboard", "Модуль главного экрана, дневного прогресса, аналитики и быстрых действий."),
        ("DashboardScreen", "Главный экран приложения с показателями дня, быстрыми действиями и переходами к разделам."),
        ("dashboard_providers.dart", "Загружает профиль, тренировки, питание, шаги, историю веса и тренерские назначения."),
        ("_saveDailyProfileMetricsSnapshot()", "Сохраняет дневные метрики пользователя в users/{userId}/daily_metrics."),
        ("DashboardAnalyticsCalculator", "Рассчитывает динамику веса, калорий, тренировок, шагов и выполнение целей."),
        ("daily_habits_widgets.dart", "Содержит виджеты дневных привычек, прогресса и плановых тренировок."),
        ("features/nutrition", "Модуль дневника питания, добавления продуктов, поиска по штрихкоду и расчета БЖУ."),
        ("AddFoodController", "Контроллер состояния формы добавления питания."),
        ("AddFoodController.setInputMethod()", "Переключает способ ввода записи питания: вручную или по штрихкоду."),
        ("AddFoodController.findProductByBarcode()", "Ищет продукт по штрихкоду в доступном каталоге."),
        ("AddFoodController.calculateMacros()", "Рассчитывает калории, белки, жиры и углеводы для выбранной порции."),
        ("AddFoodController.addFoodEntry()", "Добавляет запись в дневник питания и инициирует сохранение."),
        ("FoodDiaryController", "Контроллер дневника питания за выбранную дату."),
        ("FoodDiaryController.loadDailyFoodEntries()", "Загружает записи питания пользователя за указанную дату."),
        ("AddFoodScreen", "Экран добавления продукта в дневник питания."),
        ("FoodDiaryScreen", "Экран просмотра дневника питания по приемам пищи."),
        ("ProductDetailsScreen", "Экран отображения карточки продукта и его пищевой ценности."),
        ("SqfliteNutritionLocalDataSource", "Локальный источник данных питания, использующий базу liga_gym_nutrition.db."),
        ("SqfliteNutritionLocalDataSource.saveFoodEntry()", "Сохраняет запись питания в локальную таблицу food_entries."),
        ("SqfliteNutritionLocalDataSource.loadDailyFoodEntries()", "Загружает локальные записи питания пользователя за дату."),
        ("SqfliteNutritionLocalDataSource.markFoodEntrySynced()", "Отмечает локальную запись как синхронизированную."),
        ("FirestoreNutritionRemoteDataSource", "Удаленный источник данных питания, работающий с Cloud Firestore."),
        ("FirestoreNutritionRemoteDataSource.saveFoodEntry()", "Сохраняет запись питания в users/{userId}/food_entries."),
        ("FirestoreNutritionRemoteDataSource.saveSavedProduct()", "Сохраняет продукт быстрого доступа в saved_food_products."),
        ("InMemoryProductCatalogDataSource.findByBarcode()", "Выполняет поиск продукта по штрихкоду во встроенном каталоге."),
        ("NutritionRepositoryImpl", "Репозиторий питания, объединяющий локальный источник, Firestore и каталог продуктов."),
        ("NutritionMacroCalculator.calculate()", "Рассчитывает пищевую ценность порции на основе массы и БЖУ на 100 г."),
        ("NutritionOfflineSyncService.syncDataWithServer()", "Синхронизирует локальные записи питания с удаленным хранилищем."),
        ("features/workout", "Модуль тренировок, активной сессии, маршрута, истории и результата."),
        ("WorkoutSessionController", "Контроллер активной тренировочной сессии."),
        ("WorkoutSessionController.startWorkoutTimer()", "Запускает тренировку, таймер и сбор маршрута."),
        ("WorkoutSessionController.pauseWorkout()", "Приостанавливает тренировочную сессию."),
        ("WorkoutSessionController.resumeWorkout()", "Возобновляет приостановленную тренировку."),
        ("WorkoutSessionController.stopWorkout()", "Завершает тренировочную сессию и формирует итоговый результат."),
        ("WorkoutSessionController.saveWorkoutToDatabase()", "Сохраняет завершенную тренировку в локальное и/или облачное хранилище."),
        ("WorkoutSessionController.reset()", "Сбрасывает состояние тренировочной сессии."),
        ("WorkoutListController", "Контроллер списка и истории тренировок пользователя."),
        ("WorkoutListController.loadUserWorkouts()", "Загружает историю тренировок пользователя."),
        ("WorkoutListController.filterWorkouts()", "Фильтрует тренировки по выбранным параметрам."),
        ("StartWorkoutScreen", "Экран выбора типа тренировки и запуска сессии."),
        ("ActiveWorkoutScreen", "Экран активной тренировки с таймером, калориями, дистанцией и маршрутом."),
        ("WorkoutResultScreen", "Экран результата завершенной тренировки."),
        ("WorkoutListScreen", "Экран истории тренировок пользователя."),
        ("WorkoutRouteMap", "Виджет отображения маршрута тренировки на карте."),
        ("SqfliteWorkoutLocalDataSource", "Локальный источник данных тренировок, использующий базу liga_gym_workouts.db."),
        ("SqfliteWorkoutLocalDataSource.saveWorkout()", "Сохраняет тренировку в локальную SQLite-таблицу workouts."),
        ("SqfliteWorkoutLocalDataSource.loadUserWorkouts()", "Загружает локальную историю тренировок пользователя."),
        ("SqfliteWorkoutLocalDataSource.markWorkoutSynced()", "Отмечает тренировку как синхронизированную."),
        ("FirestoreWorkoutRemoteDataSource", "Удаленный источник данных тренировок, работающий с users/{userId}/workouts."),
        ("FirestoreWorkoutRemoteDataSource.saveWorkout()", "Сохраняет тренировку в Firestore и обновляет социальные показатели."),
        ("GeolocatorWorkoutLocationDataSource.prepareTracking()", "Проверяет разрешения и готовит получение геолокации."),
        ("GeolocatorWorkoutLocationDataSource.watchRoute()", "Передает поток точек маршрута активной тренировки."),
        ("WorkoutRepositoryImpl", "Репозиторий тренировок, объединяющий локальные, удаленные и геолокационные источники данных."),
        ("features/steps", "Модуль шагомера и дневной шаговой активности."),
        ("StepCounterScreen", "Экран просмотра текущих шагов, цели и прогресса."),
        ("StepSettingsScreen", "Экран настройки цели по шагам."),
        ("SqfliteStepLocalDataSource", "Локальный источник шагов, использующий базу liga_gym_steps.db."),
        ("SqfliteStepLocalDataSource.recordSensorReading()", "Сохраняет новое значение датчика шагов и обновляет дневной итог."),
        ("SqfliteStepLocalDataSource.loadStepCounts()", "Загружает шаги пользователя за выбранный период."),
        ("SqfliteStepLocalDataSource.loadStepsForDate()", "Возвращает количество шагов за конкретную дату."),
        ("features/water_tracker", "Модуль учета потребления воды и отображения прогресса по дневной норме."),
        ("features/workout_completion", "Модуль завершения тренировок и отображения итогового состояния после выполнения активности."),
        ("features/notifications", "Модуль экранов и логики пользовательских уведомлений."),
        ("features/social", "Социальный модуль: друзья, заявки, лидерборд, приватность и чаты."),
        ("SocialRemoteDataSource", "Источник социальных данных, работающий с коллекциями users, leaderboard_entries, friend_requests, friend_invites и interest_chats."),
        ("ensureSocialProfile()", "Создает или обновляет социальный профиль пользователя и запись лидерборда."),
        ("createFriendInvite()", "Создает код приглашения в друзья."),
        ("sendFriendRequest()", "Отправляет заявку в друзья по коду приглашения."),
        ("acceptFriendRequest()", "Принимает заявку и создает записи друзей у обоих пользователей."),
        ("removeFriend()", "Удаляет пользователя из списка друзей."),
        ("updatePrivacySettings()", "Обновляет настройки приватности пользователя."),
        ("createInterestChat()", "Создает чат по интересам или личный чат с другом."),
        ("sendChatMessage()", "Отправляет сообщение в чат и обновляет состояние комнаты."),
        ("deleteChatMessage()", "Удаляет сообщение из чата."),
        ("LeaderboardScreen", "Экран рейтинга пользователей по социальным баллам, тренировкам, калориям и шагам."),
        ("FriendRequestsScreen", "Экран входящих и исходящих заявок в друзья."),
        ("FriendsScreen", "Экран списка друзей и доступных действий."),
        ("ChatRoomScreen", "Экран переписки в выбранном чате."),
        ("TrainerMaterialsScreen", "Экран просмотра назначенных тренером материалов ученика."),
        ("features/coach", "Тренерский модуль для пользователей с ролью trainer."),
        ("CoachDashboardScreen", "Экран тренера с вкладками учеников, упражнений, рецептов и шаблонов тренировок."),
        ("CoachDashboardScreen._sendStudentInvite()", "Отправляет ученику запрос на подключение тренера."),
        ("CoachDashboardScreen._upsertExercise()", "Создает или обновляет упражнение в библиотеке тренера."),
        ("CoachDashboardScreen._deleteExercise()", "Удаляет упражнение из библиотеки тренера."),
        ("CoachDashboardScreen._upsertRecipe()", "Создает или обновляет рецепт тренера."),
        ("CoachDashboardScreen._deleteRecipe()", "Удаляет рецепт тренера."),
        ("CoachDashboardScreen._upsertTemplate()", "Создает или обновляет шаблон тренировки."),
        ("CoachDashboardScreen._deleteTemplate()", "Удаляет шаблон тренировки."),
        ("CoachDashboardScreen._assignRecipe()", "Назначает рецепт выбранному ученику."),
        ("CoachDashboardScreen._assignWorkout()", "Назначает тренировку выбранному ученику на дату и время."),
        ("_uploadCoachMedia()", "Загружает изображение или видео тренерского материала в Supabase Storage."),
        ("CoachRepositoryImpl", "Репозиторий тренерского модуля, работающий с Firestore и тренерскими коллекциями."),
        ("CoachRepositoryImpl.loadStudents()", "Загружает список активных учеников тренера."),
        ("CoachRepositoryImpl.loadLinkedTrainers()", "Загружает список тренеров, подключенных к ученику."),
        ("CoachRepositoryImpl.sendCoachRequest()", "Создает запрос тренера ученику по коду друга."),
        ("CoachRepositoryImpl.acceptCoachRequest()", "Создает активную связь ученика с тренером."),
        ("CoachRepositoryImpl.loadExercises()", "Загружает упражнения тренера."),
        ("CoachRepositoryImpl.saveExercise()", "Сохраняет упражнение тренера в Firestore."),
        ("CoachRepositoryImpl.loadRecipes()", "Загружает рецепты тренера."),
        ("CoachRepositoryImpl.saveRecipe()", "Сохраняет рецепт тренера в Firestore."),
        ("CoachRepositoryImpl.loadWorkoutTemplates()", "Загружает шаблоны тренировок тренера."),
        ("CoachRepositoryImpl.saveWorkoutTemplate()", "Сохраняет шаблон тренировки тренера."),
        ("CoachRepositoryImpl.assignRecipe()", "Создает назначение рецепта ученику."),
        ("CoachRepositoryImpl.assignWorkout()", "Создает назначение тренировки ученику."),
        ("features/ai_coach", "Модуль AI-помощника по тренировкам и питанию."),
        ("AiCoachScreen", "Экран диалога пользователя с AI-помощником."),
        ("AiCoachScreen._sendMessage()", "Формирует запрос к AI-помощнику и добавляет ответ в историю диалога."),
        ("AiCoachScreen._loadChatPages()", "Загружает сохраненные страницы чата."),
        ("AiCoachScreen._createNewPage()", "Создает новую страницу диалога."),
        ("AiCoachScreen._persistCurrentPage()", "Сохраняет текущую страницу чата локально."),
        ("FitnessAiAgent.answer()", "Формирует ответ на вопрос пользователя с учетом фитнес-контекста."),
        ("AiCompletionClient.complete()", "Отправляет сообщения в AI API и возвращает текстовый ответ."),
        ("features/exercises", "Модуль библиотеки упражнений и справочной информации по тренировкам."),
        ("domain/entities", "Сущности предметной области: пользователь, тренировка, питание, шаги, друзья, тренерские материалы."),
        ("domain/usecases", "Сценарии прикладной логики, реализующие отдельные бизнес-операции приложения."),
        ("data/models", "Модели сериализации и преобразования данных между Dart-объектами, Firestore и SQLite."),
        ("data/repositories", "Репозитории, координирующие обмен данными между presentation, domain и data слоями."),
        ("presentation/providers", "Riverpod-провайдеры зависимостей, состояний и потоков данных."),
        ("presentation/screens", "Экраны пользовательского интерфейса приложения."),
        ("presentation/widgets", "Переиспользуемые виджеты интерфейса, применяемые в нескольких сценариях."),
    ]

    add_table(doc, "Таблица 2.1.1 - Описание всех модулей и процедур", rows)
    doc.save(OUT)


if __name__ == "__main__":
    build()
