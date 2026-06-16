from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("docs/opisanie_struktury_bd_liga_gym.docx")
IMG = Path(".codex_tmp/liga_gym_db_schema.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Caption"]
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc, caption, headers, rows, widths=(2300, 1900, 1500, 3660)):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            paragraph = cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(text))
            run.font.size = Pt(9)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_metadata_table(doc):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Приложение", "Liga Gym"),
        ("Тип приложения", "мобильное приложение на Flutter"),
        ("Основные хранилища", "Firebase Firestore, локальные базы SQLite, Supabase Storage"),
        ("Аутентификация", "Firebase Authentication"),
        ("Назначение раздела", "описание структуры хранимых данных приложения"),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], "E8EEF5")
        for cell in cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_geometry(table, (2700, 6660))
    doc.add_paragraph()


def create_schema_image():
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return None

    width, height = 1500, 820
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    try:
        font_title = ImageFont.truetype("arial.ttf", 42)
        font_head = ImageFont.truetype("arial.ttf", 30)
        font_body = ImageFont.truetype("arial.ttf", 24)
    except Exception:
        font_title = ImageFont.load_default()
        font_head = ImageFont.load_default()
        font_body = ImageFont.load_default()

    draw.text((70, 50), "Схема структуры данных Liga Gym", fill="#0B2545", font=font_title)

    boxes = [
        (70, 170, 450, 380, "Firebase Firestore", ["users", "food_entries", "workouts", "social", "coach"]),
        (560, 170, 940, 380, "SQLite", ["liga_gym_nutrition.db", "liga_gym_workouts.db", "liga_gym_steps.db"]),
        (1050, 170, 1430, 380, "Supabase Storage", ["coach-media", "изображения", "видео"]),
        (315, 510, 1185, 695, "Клиентское приложение Flutter", ["синхронизация профиля, питания, тренировок, шагов и тренерских материалов"]),
    ]

    for x1, y1, x2, y2, title, lines in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, outline="#8094AA", width=4, fill="#F4F6F9")
        draw.text((x1 + 28, y1 + 26), title, fill="#1F3A5F", font=font_head)
        y = y1 + 82
        for line in lines:
            draw.text((x1 + 28, y), line, fill="#24364A", font=font_body)
            y += 38

    def arrow(start, end):
        draw.line((start, end), fill="#2E74B5", width=5)
        ex, ey = end
        sx, sy = start
        if ey > sy:
            pts = [(ex, ey), (ex - 16, ey - 24), (ex + 16, ey - 24)]
        else:
            pts = [(ex, ey), (ex - 16, ey + 24), (ex + 16, ey + 24)]
        draw.polygon(pts, fill="#2E74B5")

    arrow((260, 510), (260, 385))
    arrow((750, 510), (750, 385))
    arrow((1240, 510), (1240, 385))

    IMG.parent.mkdir(parents=True, exist_ok=True)
    img.save(IMG)
    return IMG


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    caption.font.size = Pt(10)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "Liga Gym - структура базы данных"
    header.style = doc.styles["Normal"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Раздел 1.6")

    doc.add_heading("1.6 Описание структуры базы данных", level=1)
    doc.add_paragraph(
        "Схема данных - это структура базы данных, описанная на формальном "
        "языке, поддерживаемом используемыми средствами хранения. Для приложения "
        "Liga Gym схема включает облачные коллекции Firebase Firestore, локальные "
        "таблицы SQLite и объектное хранилище Supabase Storage для медиафайлов."
    )
    doc.add_paragraph(
        "Схема отношений базы данных представлена в приложении А. При проектировании "
        "структуры хранения данных использовалась гибридная модель: Firestore "
        "применяется для синхронизируемых пользовательских, социальных и тренерских "
        "данных; SQLite обеспечивает локальную автономную работу дневника питания, "
        "тренировок и шаговой активности; Supabase Storage хранит изображения и "
        "видео тренерских материалов."
    )
    doc.add_paragraph(
        "Для управления учетными записями используется Firebase Authentication. "
        "Данные аутентификации хранятся в отдельном сервисе Firebase и не входят "
        "в состав описываемых коллекций Firestore и таблиц SQLite. Описание "
        "структуры базы данных приведено в таблицах 1.6.1 - 1.6.15."
    )

    add_metadata_table(doc)

    headers = ["Содержание поля", "Имя поля", "Тип, длина", "Примечания"]

    add_table(doc, "Таблица 1.6.1 - users (Пользователи)", headers, [
        ("Идентификатор пользователя", "userId", "String", "Ключ документа, соответствует Firebase Auth uid"),
        ("Адрес электронной почты", "email", "String", "Используется при входе и отображении профиля"),
        ("Имя", "name", "String", "Отображаемое имя пользователя"),
        ("Пол", "gender", "String", "male / female"),
        ("Дата рождения", "birthDate", "Timestamp", "Дата рождения пользователя"),
        ("Город", "city", "String", "Используется для городского лидерборда"),
        ("Код друга", "friendCode", "String", "Код приглашения в друзья"),
        ("Код страны", "countryCode", "String", "По умолчанию RU"),
        ("Рост", "heightCm", "Number", "Рост пользователя в сантиметрах"),
        ("Начальный вес", "startWeightKg", "Number", "Исходная масса тела"),
        ("Текущий вес", "currentWeightKg", "Number", "Актуальный показатель веса"),
        ("Целевой вес", "targetWeightKg", "Number", "Используется в аналитике веса"),
        ("Тип цели", "goalType", "String", "loseWeight / maintainWeight / gainWeight"),
        ("Цель по шагам", "dailyStepGoal", "Number", "По умолчанию 10000"),
        ("Цель по калориям", "dailyCalorieGoal", "Number", "По умолчанию 2200"),
        ("Роль", "role", "String", "student / trainer"),
        ("Социальный рейтинг", "socialScore", "Number", "Суммарные баллы активности"),
        ("Количество тренировок", "socialWorkoutsCount", "Number", "Показатель для социальных функций"),
        ("Сожженные калории", "socialCaloriesBurned", "Number", "Показатель для рейтинга"),
        ("Количество шагов", "socialStepsCount", "Number", "Показатель для рейтинга"),
        ("Видимость в лидерборде", "visibleInFriendLeaderboard", "Boolean", "Настройка приватности рейтинга"),
        ("Дата создания", "createdAt", "Timestamp", "Серверное время создания"),
        ("Дата обновления", "updatedAt", "Timestamp", "Серверное время обновления"),
    ])

    add_table(doc, "Таблица 1.6.2 - users/{userId}/weight_history (История веса)", headers, [
        ("Дата записи", "recordedAt", "Timestamp", "Время фиксации веса"),
        ("Вес", "weightKg", "Number", "Масса тела в килограммах"),
        ("Дата обновления", "updatedAt", "Timestamp", "Служебная отметка изменения"),
    ])

    add_table(doc, "Таблица 1.6.3 - users/{userId}/food_entries (Записи дневника питания)", headers, [
        ("Идентификатор записи", "id", "String", "Идентификатор документа"),
        ("Тип приема пищи", "meal_type", "String", "breakfast / lunch / dinner / snack"),
        ("Ключ даты", "date_key", "String", "Формат YYYY-MM-DD"),
        ("Название продукта (EN)", "product_name_en", "String", "Английское название продукта"),
        ("Название продукта (RU)", "product_name_ru", "String", "Русское название продукта"),
        ("Штрихкод", "barcode", "String", "Необязательное поле"),
        ("Вес порции", "grams", "Number", "Размер порции в граммах"),
        ("Калорийность", "calories", "Number", "Расчетная калорийность"),
        ("Белки", "proteins", "Number", "Количество белков"),
        ("Жиры", "fats", "Number", "Количество жиров"),
        ("Углеводы", "carbs", "Number", "Количество углеводов"),
        ("Дата и время записи", "logged_at", "Timestamp", "Время добавления записи"),
        ("Способ ввода", "input_method", "String", "manual / barcode"),
    ])

    add_table(doc, "Таблица 1.6.4 - users/{userId}/saved_food_products (Сохраненные продукты)", headers, [
        ("Идентификатор продукта", "id", "String", "Идентификатор документа"),
        ("Название продукта (EN)", "product_name_en", "String", "Английское название"),
        ("Название продукта (RU)", "product_name_ru", "String", "Русское название"),
        ("Штрихкод", "barcode", "String", "Необязательный штрихкод"),
        ("Калорийность", "calories", "Number", "Значение на 100 г"),
        ("Белки", "proteins", "Number", "Значение на 100 г"),
        ("Жиры", "fats", "Number", "Значение на 100 г"),
        ("Углеводы", "carbs", "Number", "Значение на 100 г"),
    ])

    add_table(doc, "Таблица 1.6.5 - users/{userId}/workouts (Тренировки)", headers, [
        ("Идентификатор тренировки", "id", "String", "Идентификатор документа"),
        ("Пользователь", "userId", "String", "Идентификатор владельца тренировки"),
        ("Тип тренировки", "type", "String", "running / cycling / walking / strength / cardio"),
        ("Дата и время начала", "startedAt", "String", "Формат ISO 8601"),
        ("Дата и время окончания", "endedAt", "String", "Формат ISO 8601"),
        ("Продолжительность", "durationSeconds", "Number", "В секундах"),
        ("Калории", "calories", "Number", "Расчетный расход калорий"),
        ("Дистанция", "distanceMeters", "Number", "В метрах"),
        ("Маршрут", "routeJson", "String", "JSON-массив точек маршрута"),
        ("Название", "title", "String", "Пользовательское название тренировки"),
        ("Заметка", "note", "String", "Комментарий пользователя"),
        ("Место", "place", "String", "Место выполнения"),
        ("Упражнения", "exerciseEntriesJson", "String", "JSON-массив упражнений"),
        ("Ручной ввод", "isManual", "Boolean", "Признак ручного добавления"),
        ("Признак синхронизации", "isSynced", "Boolean", "Для облачной записи принимает true"),
    ])

    add_table(doc, "Таблица 1.6.6 - users/{userId}/daily_metrics (Дневные показатели)", headers, [
        ("Ключ даты", "dateKey", "String", "Идентификатор документа за день"),
        ("Дата", "date", "Timestamp", "Дата формирования показателей"),
        ("Шаги", "steps", "Number", "Количество шагов за день"),
        ("Наличие шагов", "hasRecordedSteps", "Boolean", "Признак записанных шагов"),
        ("Потребленные калории", "caloriesConsumed", "Number", "Итог по дневнику питания"),
        ("Сожженные калории", "caloriesBurned", "Number", "Итог по тренировкам"),
        ("Белки", "proteins", "Number", "Итог за день"),
        ("Жиры", "fats", "Number", "Итог за день"),
        ("Углеводы", "carbs", "Number", "Итог за день"),
        ("Количество записей питания", "foodEntriesCount", "Number", "Количество приемов пищи"),
        ("Количество тренировок", "workoutsCount", "Number", "Количество тренировок за день"),
        ("Минуты тренировок", "workoutMinutes", "Number", "Суммарная длительность"),
        ("Дистанция тренировок", "workoutDistanceMeters", "Number", "Суммарная дистанция"),
        ("Цель по шагам", "stepGoal", "Number", "Пользовательская цель"),
        ("Цель по калориям", "calorieGoal", "Number", "Пользовательская цель"),
        ("Дата обновления", "updatedAt", "Timestamp", "Серверное время обновления"),
    ])

    add_table(doc, "Таблица 1.6.7 - liga_gym_nutrition.db / food_entries (Локальные записи питания)", headers, [
        ("Идентификатор записи", "id", "TEXT", "Первичный ключ"),
        ("Пользователь", "user_id", "TEXT", "Обязательное поле"),
        ("Тип приема пищи", "meal_type", "TEXT", "Обязательное поле"),
        ("Ключ даты", "date_key", "TEXT", "Используется для выборки по дате"),
        ("Название продукта (EN)", "product_name_en", "TEXT", "Обязательное поле"),
        ("Название продукта (RU)", "product_name_ru", "TEXT", "Обязательное поле"),
        ("Штрихкод", "barcode", "TEXT", "Необязательное поле"),
        ("Вес порции", "grams", "REAL", "Обязательное поле"),
        ("Калорийность", "calories", "REAL", "Обязательное поле"),
        ("Белки", "proteins", "REAL", "Обязательное поле"),
        ("Жиры", "fats", "REAL", "Обязательное поле"),
        ("Углеводы", "carbs", "REAL", "Обязательное поле"),
        ("Дата и время записи", "logged_at", "INTEGER", "Unix time в миллисекундах"),
        ("Способ ввода", "input_method", "TEXT", "manual / barcode"),
        ("Признак синхронизации", "is_synced", "INTEGER", "0 - не синхронизировано, 1 - синхронизировано"),
    ])

    add_table(doc, "Таблица 1.6.8 - liga_gym_workouts.db / workouts (Локальные тренировки)", headers, [
        ("Идентификатор тренировки", "id", "TEXT", "Первичный ключ"),
        ("Пользователь", "user_id", "TEXT", "Обязательное поле"),
        ("Тип тренировки", "type", "TEXT", "Обязательное поле"),
        ("Дата и время начала", "started_at", "INTEGER", "Unix time в миллисекундах"),
        ("Дата и время окончания", "ended_at", "INTEGER", "Unix time в миллисекундах"),
        ("Продолжительность", "duration_seconds", "INTEGER", "В секундах"),
        ("Калории", "calories", "REAL", "Обязательное поле"),
        ("Дистанция", "distance_meters", "REAL", "В метрах"),
        ("Маршрут", "route_json", "TEXT", "JSON-массив точек маршрута"),
        ("Признак синхронизации", "is_synced", "INTEGER", "0 - не синхронизировано, 1 - синхронизировано"),
        ("Название", "title", "TEXT", "Необязательное поле"),
        ("Заметка", "note", "TEXT", "Необязательное поле"),
        ("Место", "place", "TEXT", "Необязательное поле"),
        ("Упражнения", "exercise_entries_json", "TEXT", "JSON-массив упражнений"),
        ("Ручной ввод", "is_manual", "INTEGER", "0 - нет, 1 - да"),
    ])

    add_table(doc, "Таблица 1.6.9 - liga_gym_steps.db / daily_steps (Локальный учет шагов)", headers, [
        ("Пользователь", "user_id", "TEXT", "Составной первичный ключ"),
        ("Ключ даты", "date_key", "TEXT", "Составной первичный ключ, формат YYYY-MM-DD"),
        ("Количество шагов", "steps", "INTEGER", "Итоговое значение за день"),
        ("Дата обновления", "updated_at", "INTEGER", "Unix time в миллисекундах"),
    ])

    add_table(doc, "Таблица 1.6.10 - liga_gym_steps.db / step_sensor_state (Состояние датчика шагов)", headers, [
        ("Пользователь", "user_id", "TEXT", "Первичный ключ"),
        ("Последнее значение датчика", "last_sensor_steps", "INTEGER", "Последнее считанное значение"),
        ("Время последнего значения", "last_sensor_timestamp", "INTEGER", "Unix time в миллисекундах"),
    ])

    add_table(doc, "Таблица 1.6.11 - leaderboard_entries (Лидерборд)", headers, [
        ("Идентификатор пользователя", "userId", "String", "Ключ документа"),
        ("Отображаемое имя", "displayName", "String", "Имя в рейтинге"),
        ("Город", "city", "String", "Используется для городского рейтинга"),
        ("Очки", "score", "Number", "Суммарный социальный рейтинг"),
        ("Количество тренировок", "workoutsCount", "Number", "Показатель активности"),
        ("Сожженные калории", "caloriesBurned", "Number", "Показатель активности"),
        ("Количество шагов", "stepsCount", "Number", "Показатель активности"),
        ("Видимость", "visibleInFriendLeaderboard", "Boolean", "Разрешение показа в рейтинге"),
        ("Дата обновления", "updatedAt", "Timestamp", "Серверное время обновления"),
    ])

    add_table(doc, "Таблица 1.6.12 - friend_invites и friend_requests (Приглашения и заявки в друзья)", headers, [
        ("Код приглашения", "friend_invites/{inviteId}", "Document", "inviteId соответствует friendCode"),
        ("Владелец кода", "ownerUserId", "String", "Идентификатор пользователя"),
        ("Имя владельца", "ownerDisplayName", "String", "Имя отправителя приглашения"),
        ("Почта владельца", "ownerEmail", "String", "Почта отправителя приглашения"),
        ("Заявка в друзья", "friend_requests/{requestId}", "Document", "Документ запроса на добавление"),
        ("Отправитель", "fromUserId", "String", "Пользователь, отправивший заявку"),
        ("Получатель", "toUserId", "String", "Пользователь, получивший заявку"),
        ("Статус", "status", "String", "pending / accepted / declined"),
        ("Дата создания", "createdAt", "Timestamp", "Серверное время создания"),
        ("Дата обновления", "updatedAt", "Timestamp", "Серверное время обновления"),
    ])

    add_table(doc, "Таблица 1.6.13 - interest_chats (Чаты по интересам)", headers, [
        ("Идентификатор чата", "chatId", "String", "Ключ документа"),
        ("Название", "title", "String", "Название комнаты"),
        ("Описание", "description", "String", "Описание чата"),
        ("Тип", "type", "String", "friend_dm или другой тип комнаты"),
        ("Участники", "participantIds", "Array<String>", "Список пользователей"),
        ("Создатель", "createdBy", "String", "Идентификатор создателя"),
        ("Количество участников", "memberCount", "Number", "Счетчик участников"),
        ("Участники", "members/{userId}", "Subcollection", "Профиль участника чата"),
        ("Сообщения", "messages/{messageId}", "Subcollection", "senderId, message, sentAt, type"),
        ("Журнал модерации", "moderation_logs/{logId}", "Subcollection", "Служебные события модерации"),
    ])

    add_table(doc, "Таблица 1.6.14 - Тренерские коллекции Firestore", headers, [
        ("Запрос тренеру", "coach_requests", "Collection", "trainerId, studentId, status, createdAt, updatedAt"),
        ("Связь ученика с тренером", "users/{studentId}/coach_links", "Subcollection", "trainerId, studentId, status, linkedAt"),
        ("Тренер", "trainers/{trainerId}", "Document", "Корневой документ материалов тренера"),
        ("Упражнения", "trainers/{trainerId}/exercises", "Subcollection", "title, description, videoUrl, media, muscleGroups, equipment, techniqueText"),
        ("Рецепты", "trainers/{trainerId}/recipes", "Subcollection", "title, description, ingredientsText, macros, media, servingGrams"),
        ("Шаблоны тренировок", "trainers/{trainerId}/workout_templates", "Subcollection", "title, goal, instructions, exerciseIds"),
        ("Назначенные рецепты", "users/{studentId}/recipe_assignments", "Subcollection", "recipeId, trainerId, trainerName, status, createdAt"),
        ("Назначенные тренировки", "users/{studentId}/workout_assignments", "Subcollection", "templateId, trainerId, scheduledAt, status"),
    ])

    add_table(doc, "Таблица 1.6.15 - Supabase Storage (Медиафайлы тренера)", headers, [
        ("Хранилище", "coach-media", "Bucket", "Объектное хранилище Supabase Storage"),
        ("Путь файла", "trainers/{trainerId}/{type}/{timestamp}", "String", "Структура пути загружаемых файлов"),
        ("URL файла", "url", "String", "Ссылка сохраняется в поле media коллекций тренера"),
        ("Имя файла", "name", "String", "Отображаемое имя вложения"),
        ("Тип файла", "type", "String", "image / video"),
    ])

    image_path = create_schema_image()
    if image_path is not None:
        doc.add_paragraph("На рисунке 1.6 изображена схема хранения данных приложения.")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(15.2))
        cap = doc.add_paragraph("Рисунок 1.6 - Схема структуры данных приложения")
        cap.style = doc.styles["Caption"]
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Схема базы данных представляет собой формализованное описание структуры "
        "хранимой информации, включающее облачные коллекции, локальные таблицы, "
        "медиафайлы и связи между ними. В гибридной архитектуре Liga Gym она "
        "определяет состав данных, правила синхронизации и способы обеспечения "
        "целостности информации при работе приложения онлайн и офлайн."
    )
    doc.add_paragraph(
        "В связи с расширением функциональности структура хранения дополнена "
        "социальными коллекциями, дневными метриками, локальным учетом шагов и "
        "тренерским режимом. Это обеспечивает ведение профиля, дневника питания, "
        "истории тренировок, лидерборда, общения пользователей и назначения "
        "тренировочных материалов ученикам."
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
