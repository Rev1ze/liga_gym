from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"C:\Users\UFANB\liga_gym")
DOC_PATH = Path(
    r"C:\Users\UFANB\Desktop\докс\исправка\1\ne trogat\ласт диплом"
) / "Хуснияров_Усман_Диплом_Отчет_финальная_версия_по_ПЗ.docx"
OUT_PATH = Path(
    r"C:\Users\UFANB\Desktop\докс\исправка\1\ne trogat\ласт диплом"
) / "Хуснияров_Усман_Диплом_Отчет_финальная_версия_по_ПЗ_код_сокращен.docx"


FRAGMENTS = [
    (
        "А.1 Инициализация приложения",
        "lib/main.dart",
        None,
        70,
    ),
    (
        "А.2 Маршруты экранов приложения",
        "lib/core/navigation/app_routes.dart",
        None,
        45,
    ),
    (
        "А.3 Генерация экранов по маршрутам",
        "lib/core/navigation/app_router.dart",
        None,
        130,
    ),
    (
        "А.4 Модель профиля пользователя",
        "lib/features/auth/data/models/user_profile_model.dart",
        None,
        105,
    ),
    (
        "А.5 Источник данных профиля Firestore",
        "lib/features/auth/data/datasources/profile_remote_data_source.dart",
        None,
        155,
    ),
    (
        "А.6 Модель записи питания",
        "lib/features/nutrition/data/models/food_entry_model.dart",
        None,
        155,
    ),
    (
        "А.7 Локальное хранение питания SQLite",
        "lib/features/nutrition/data/datasources/nutrition_local_data_source.dart",
        None,
        125,
    ),
    (
        "А.8 Облачное хранение питания Firestore",
        "lib/features/nutrition/data/datasources/nutrition_remote_data_source.dart",
        None,
        115,
    ),
    (
        "А.9 Модель тренировки и маршрута",
        "lib/features/workout/data/models/workout_model.dart",
        None,
        175,
    ),
    (
        "А.10 Локальное хранение тренировок SQLite",
        "lib/features/workout/data/datasources/workout_local_data_source.dart",
        None,
        150,
    ),
    (
        "А.11 Локальный учет шагов",
        "lib/features/steps/data/datasources/step_local_data_source.dart",
        None,
        150,
    ),
    (
        "А.12 Расчет аналитики панели управления",
        "lib/features/dashboard/domain/services/dashboard_analytics_calculator.dart",
        None,
        150,
    ),
    (
        "А.13 Социальные функции и приватность",
        "lib/features/social/data/datasources/social_remote_data_source.dart",
        340,
        250,
    ),
    (
        "А.14 Репозиторий тренера и материалов",
        "lib/features/coach/data/repositories/coach_repository_impl.dart",
        None,
        220,
    ),
    (
        "А.15 AI-помощник Liga Gym",
        "lib/features/ai_coach/domain/services/fitness_ai_agent.dart",
        None,
        215,
    ),
    (
        "А.16 Базовый сервис синхронизации",
        "lib/core/offline/offline_sync_service.dart",
        None,
        35,
    ),
]


def paragraph_text(paragraph) -> str:
    return " ".join(paragraph.text.split())


def delete_between(doc: Document, start_idx: int, end_idx: int) -> None:
    for paragraph in doc.paragraphs[start_idx:end_idx]:
        paragraph._element.getparent().remove(paragraph._element)


def insert_before(anchor, paragraph):
    anchor._p.addprevious(paragraph._p)


def add_before(doc: Document, anchor, text: str = "", *, kind: str = "body"):
    p = doc.add_paragraph()
    insert_before(anchor, p)
    if text:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if kind == "code":
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(8.5)
        else:
            run.font.size = Pt(14)
            run.bold = kind in {"title", "heading"}
    p.paragraph_format.space_after = Pt(0)
    if kind == "title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.5
    elif kind == "heading":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.5
    elif kind == "code":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
    return p


def read_fragment(relative_path: str, start_line: int | None, limit: int) -> list[str]:
    path = ROOT / relative_path
    lines = path.read_text(encoding="utf-8").splitlines()
    if start_line is None:
        start = 0
    else:
        start = max(0, start_line - 1)
    fragment = lines[start : start + limit]
    return [line.rstrip() for line in fragment]


def main() -> None:
    doc = Document(str(DOC_PATH))
    app_a_idx = next(
        i for i, p in enumerate(doc.paragraphs) if paragraph_text(p) == "ПРИЛОЖЕНИЕ А"
    )
    app_d_idx = next(
        i for i, p in enumerate(doc.paragraphs) if paragraph_text(p) == "ПРИЛОЖЕНИЕ Д"
    )

    delete_between(doc, app_a_idx, app_d_idx)
    anchor = next(p for p in doc.paragraphs if paragraph_text(p) == "ПРИЛОЖЕНИЕ Д")

    page = add_before(doc, anchor)
    page.add_run().add_break(WD_BREAK.PAGE)
    add_before(doc, anchor, "ПРИЛОЖЕНИЕ А", kind="title")
    add_before(doc, anchor, "Исходный код основных модулей программы", kind="title")
    add_before(
        doc,
        anchor,
        "В приложении приведены сокращенные фрагменты исходного кода мобильного приложения Liga Gym. Отобраны файлы, подтверждающие реализацию инициализации, маршрутизации, профиля, тренировок, питания, шагомера, аналитики, социальных функций, тренерского режима, AI-помощника и синхронизации данных.",
    )

    for title, rel_path, start_line, limit in FRAGMENTS:
        add_before(doc, anchor, title, kind="heading")
        add_before(doc, anchor, f"Файл: {rel_path}", kind="heading")
        for line in read_fragment(rel_path, start_line, limit):
            add_before(doc, anchor, line if line else " ", kind="code")
        add_before(doc, anchor, " ")

    page = add_before(doc, anchor)
    page.add_run().add_break(WD_BREAK.PAGE)

    doc.save(str(OUT_PATH))
    print(OUT_PATH)


if __name__ == "__main__":
    main()
