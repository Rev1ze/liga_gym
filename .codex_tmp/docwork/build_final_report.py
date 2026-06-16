from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


WORK = Path(__file__).resolve().parent
SRC = WORK / "source_husniyarov.docx"
OUT = WORK / "final_working.docx"


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def style_table(table, widths: list[float] | None = None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table_font_size = Pt(10 if len(table.columns) >= 5 else 11)
    set_borders(table)
    if table.rows:
        set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths and c_idx < len(widths):
                cell.width = Cm(widths[c_idx])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = table_font_size
                    if r_idx == 0:
                        run.bold = True
                if r_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx == 0:
                shade_cell(cell, "D9EAF7")


def move_after(anchor, element):
    if hasattr(anchor, "_p"):
        anchor._p.addnext(element)
    elif hasattr(anchor, "_tbl"):
        anchor._tbl.addnext(element)
    else:
        anchor.addnext(element)


def add_paragraph_after(doc: Document, anchor, text: str = "", *, bold=False, align=None):
    p = doc.add_paragraph()
    move_after(anchor, p._p)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_caption(doc: Document, anchor, text: str):
    p = add_paragraph_after(doc, anchor, text, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        run.font.size = Pt(12)
        run.italic = True
    return p


def add_table_after(doc: Document, anchor, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    move_after(anchor, table._tbl)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], str(value))
    style_table(table, widths)
    return table


def find_paragraph(doc: Document, startswith: str):
    for p in doc.paragraphs:
        if " ".join(p.text.split()).startswith(startswith):
            return p
    raise ValueError(f"Paragraph not found: {startswith}")


def normalize_document(doc: Document):
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    for p in doc.paragraphs:
        text = " ".join(p.text.split())
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(14)
        if text.startswith(("ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "СОДЕРЖАНИЕ", "ПРИЛОЖЕНИЕ")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            for run in p.runs:
                run.bold = True
        elif len(text) > 3 and (text[:2].isdigit() or text[:3].replace(".", "").isdigit()):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            for run in p.runs:
                run.bold = True
        elif text.startswith(("Таблица", "Рисунок", "Продолжение таблицы")):
            p.paragraph_format.first_line_indent = Cm(0)
            for run in p.runs:
                run.font.size = Pt(12)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
    for table in doc.tables:
        style_table(table)


def insert_intro(doc: Document):
    anchor = find_paragraph(doc, "ВВЕДЕНИЕ")
    paragraphs = [
        "Актуальность темы определяется тем, что современные пользователи нуждаются в единой цифровой среде для контроля физической активности, питания и динамики достижения персональных целей. Разрозненное ведение тренировок, дневника питания, шагов и веса затрудняет анализ состояния пользователя, поэтому разработка мобильной информационной системы мониторинга физической активности и питания имеет практическую значимость.",
        "Целью дипломного проекта является разработка информационной системы Liga Gym, обеспечивающей регистрацию и авторизацию пользователя, заполнение профиля, настройку целей, учет тренировок и маршрутов, ведение дневника питания, расчет калорийности и БЖУ, фиксацию шаговой активности, анализ прогресса, формирование PDF-отчета, социальное взаимодействие, работу с тренером и получение фитнес-рекомендаций через AI-помощника.",
        "Для достижения цели требуется решить следующие задачи: изучить предметную область мониторинга активности и питания; определить входные и выходные данные; спроектировать структуру хранения данных; разработать мобильное приложение на Flutter и Dart; реализовать интеграцию с Firebase Authentication, Cloud Firestore, SQLite, Supabase Storage и Riverpod; выполнить тестирование основных пользовательских сценариев.",
        "Объектом автоматизации является процесс персонального учета физической активности, питания и тренировочных результатов пользователя. Предметом разработки является мобильное приложение Liga Gym, в котором реализованы интерфейсные формы, сервисы хранения данных, локальная работа, синхронизация с облаком, аналитика и пользовательские сценарии фитнес-дневника.",
    ]
    for text in reversed(paragraphs):
        add_paragraph_after(doc, anchor, text)


def insert_subject_area(doc: Document):
    anchor = find_paragraph(doc, "1.1 Описание предметной области")
    rows = [
        ["Пользователь", "Учетная запись и профиль владельца данных", "email, name, gender, birthDate, city, role, friendCode"],
        ["Цели пользователя", "Параметры контроля прогресса", "goalType, dailyStepGoal, dailyCalorieGoal, targetWeightKg"],
        ["История веса", "Фиксация динамики массы тела", "recordedAt, weightKg"],
        ["Тренировка", "Сохраненная спортивная активность", "type, startedAt, endedAt, duration, calories, distanceMeters, route"],
        ["Точка маршрута", "Координата движения при тренировке", "latitude, longitude, recordedAt, speedMetersPerSecond"],
        ["Запись питания", "Факт приема продукта за выбранную дату", "mealType, productNameRu, grams, calories, proteins, fats, carbs"],
        ["Продукт", "Справочная или сохраненная карточка продукта", "name_ru, barcode, calories, proteins, fats, carbs"],
        ["Шаговая активность", "Итог шагов пользователя за дату", "user_id, date_key, steps, updated_at"],
        ["Друг и заявка в друзья", "Социальная связь пользователей", "fromUserId, toUserId, status, displayName"],
        ["Чат и сообщение", "Обмен сообщениями и результатами", "title, memberCount, senderId, message, sentAt, type"],
        ["Лидерборд", "Рейтинг активности с учетом приватности", "score, workoutsCount, caloriesBurned, stepsCount"],
        ["Тренерский материал", "Упражнения, рецепты и шаблоны тренировок", "title, description, media, exerciseIds, scheduledAt"],
    ]
    cap = add_caption(doc, anchor, "Таблица 1.А – Основные сущности предметной области Liga Gym")
    table = add_table_after(doc, cap, ["Сущность", "Назначение", "Основные атрибуты"], rows, [4.2, 6.6, 6.5])
    add_paragraph_after(doc, table, "Представленные сущности соответствуют фактической структуре проекта: модели профиля находятся в модуле auth, тренировки и маршруты - в модуле workout, питание - в модуле nutrition, шаги - в модуле steps, социальные функции - в модуле social, тренерские материалы - в модуле coach.")


def insert_processes(doc: Document):
    anchor = find_paragraph(doc, "1.2 Проектирование бизнес-процессов предметной области")
    rows = [
        ["Регистрация и профиль", "Пользователь вводит email, пароль и персональные данные", "Создается учетная запись Firebase Auth и документ users"],
        ["Настройка целей", "Пользователь задает шаги, калории и целевой вес", "Параметры используются на панели управления и в аналитике"],
        ["Запись тренировки", "Пользователь выбирает тип активности, запускает сессию и завершает ее", "Сохраняются длительность, калории, дистанция и маршрут"],
        ["Учет питания", "Пользователь выбирает продукт вручную или по штрихкоду", "Система рассчитывает КБЖУ по массе порции"],
        ["Шаговая активность", "Сервис получает показания датчика шагов", "Локально сохраняется дневной итог шагов"],
        ["Аналитика", "Система агрегирует питание, тренировки, шаги и вес", "Пользователь видит прогресс и может сформировать PDF-отчет"],
        ["Социальные функции", "Пользователь отправляет заявки, общается в чате и смотрит рейтинг", "Доступ к данным ограничивается настройками приватности"],
        ["Работа тренера", "Тренер создает упражнения, рецепты и шаблоны", "Материалы назначаются ученикам и доступны в приложении"],
    ]
    cap = add_caption(doc, anchor, "Таблица 1.Б – Бизнес-процессы предметной области")
    add_table_after(doc, cap, ["Процесс", "Входное действие", "Результат обработки"], rows, [4.5, 6.2, 6.6])


def insert_requirements(doc: Document):
    anchor = find_paragraph(doc, "1.5 Общие требования к программному продукту")
    rows = [
        ["Функциональные", "Регистрация, вход, профиль, цели, тренировки, питание, шаги, аналитика, PDF, AI, друзья, чат, лидерборд, тренерские материалы"],
        ["Надежность", "Обработка ошибок Firebase Auth, сетевых тайм-аутов, локальное сохранение записей и последующая синхронизация"],
        ["Данные", "Хранение пользовательских данных в Cloud Firestore, локальных записей в SQLite, настроек в Shared Preferences"],
        ["Безопасность", "Разграничение доступа правилами Firestore, учет владельца документа userId и настройка приватности социальных данных"],
        ["Интерфейс", "Русская и английская локализация, адаптивные экраны Flutter, единая навигация через AppRoutes"],
        ["Производительность", "Локальные выборки SQLite, ограничение сетевых запросов по тайм-ауту, фоновая обработка шагомера"],
        ["Переносимость", "Проект содержит конфигурации Android, iOS, Web, Windows, macOS и Linux, основная целевая среда - мобильное приложение"],
    ]
    cap = add_caption(doc, anchor, "Таблица 1.В – Общие требования к программному продукту Liga Gym")
    add_table_after(doc, cap, ["Группа требований", "Содержание требования"], rows, [4.5, 12.8])


def field_rows(fields: Iterable[tuple[str, str, str, str]]) -> list[list[str]]:
    return [[a, b, c, d] for a, b, c, d in fields]


def insert_db_tables(doc: Document):
    anchor = find_paragraph(doc, "1.6 Описание структуры базы данных")
    blocks = [
        ("Таблица 1.Г – Коллекция users", [
            ("Идентификатор пользователя", "userId", "String", "Ключ документа, соответствует Firebase Auth uid"),
            ("Электронная почта", "email", "String", "Используется при входе и отображении профиля"),
            ("Имя", "name", "String", "Отображаемое имя пользователя"),
            ("Пол", "gender", "String", "male/female из перечисления Gender"),
            ("Дата рождения", "birthDate", "Timestamp", "Используется при заполнении профиля"),
            ("Город", "city", "String", "Дополнительный параметр профиля"),
            ("Код друга", "friendCode", "String", "Используется для приглашений"),
            ("Рост", "heightCm", "Number", "Рост пользователя в сантиметрах"),
            ("Начальный вес", "startWeightKg", "Number", "Исходная масса тела"),
            ("Текущий вес", "currentWeightKg", "Number", "Актуальный показатель"),
            ("Целевой вес", "targetWeightKg", "Number", "Используется в аналитике веса"),
            ("Тип цели", "goalType", "String", "loseWeight, maintainWeight, gainWeight"),
            ("Цель по шагам", "dailyStepGoal", "Number", "По умолчанию 10000"),
            ("Цель по калориям", "dailyCalorieGoal", "Number", "По умолчанию 2200"),
            ("Роль", "role", "String", "student или trainer"),
        ]),
        ("Таблица 1.Д – Подколлекция users/{userId}/weight_history", [
            ("Дата записи", "recordedAt", "Timestamp", "Время фиксации веса"),
            ("Вес", "weightKg", "Number", "Масса тела в килограммах"),
            ("Дата обновления", "updatedAt", "Timestamp", "Служебная отметка изменения"),
        ]),
        ("Таблица 1.Е – Подколлекция users/{userId}/food_entries", [
            ("Идентификатор", "id", "String", "Ключ записи питания"),
            ("Прием пищи", "meal_type", "String", "breakfast, lunch, dinner, snack"),
            ("Дата", "date_key", "String", "Ключ дня в формате YYYY-MM-DD"),
            ("Название EN", "product_name_en", "String", "Английское название продукта"),
            ("Название RU", "product_name_ru", "String", "Русское название продукта"),
            ("Штрихкод", "barcode", "String", "Заполняется при поиске по штрихкоду"),
            ("Масса", "grams", "Number", "Размер порции в граммах"),
            ("Калории", "calories", "Number", "Расчетная калорийность порции"),
            ("Белки", "proteins", "Number", "Расчетное количество белков"),
            ("Жиры", "fats", "Number", "Расчетное количество жиров"),
            ("Углеводы", "carbs", "Number", "Расчетное количество углеводов"),
            ("Дата ввода", "logged_at", "Timestamp", "Время добавления записи"),
            ("Способ ввода", "input_method", "String", "manual или barcode"),
        ]),
        ("Таблица 1.Ж – Подколлекция users/{userId}/workouts", [
            ("Идентификатор", "id", "String", "Ключ тренировки"),
            ("Пользователь", "userId", "String", "Владелец записи"),
            ("Тип тренировки", "type", "String", "Тип из перечисления WorkoutType"),
            ("Начало", "startedAt", "String", "Дата и время начала ISO-8601"),
            ("Окончание", "endedAt", "String", "Дата и время окончания ISO-8601"),
            ("Длительность", "durationSeconds", "Number", "Продолжительность в секундах"),
            ("Калории", "calories", "Number", "Расчетный расход калорий"),
            ("Дистанция", "distanceMeters", "Number", "Расстояние в метрах"),
            ("Маршрут", "routeJson", "String", "JSON-массив точек маршрута"),
            ("Название", "title", "String", "Пользовательское название тренировки"),
            ("Заметка", "note", "String", "Комментарий пользователя"),
            ("Место", "place", "String", "Место выполнения тренировки"),
            ("Упражнения", "exerciseEntriesJson", "String", "JSON-массив упражнений"),
            ("Ручной ввод", "isManual", "Boolean", "Признак ручного добавления"),
        ]),
        ("Таблица 1.И – Локальные таблицы SQLite", [
            ("Записи питания", "liga_gym_nutrition.db / food_entries", "SQLite", "Локальный дневник питания и флаг is_synced"),
            ("Тренировки", "liga_gym_workouts.db / workouts", "SQLite", "Кэш тренировок, маршрута и упражнений"),
            ("Дневные шаги", "liga_gym_steps.db / daily_steps", "SQLite", "Итоги шагов по user_id и date_key"),
            ("Состояние датчика", "liga_gym_steps.db / step_sensor_state", "SQLite", "Последние показания датчика шагов"),
        ]),
        ("Таблица 1.К – Социальные коллекции Firestore", [
            ("Лидерборд", "leaderboard_entries", "Collection", "score, workoutsCount, caloriesBurned, stepsCount"),
            ("Заявки в друзья", "friend_requests", "Collection", "fromUserId, toUserId, status, createdAt"),
            ("Приглашения", "friend_invites", "Collection", "Код друга и владелец приглашения"),
            ("Друзья", "users/{userId}/friends", "Subcollection", "Профиль друга и доступные категории"),
            ("Настройки приватности", "users/{userId}/social_settings/privacy", "Document", "allowedCategories, видимость рейтинга"),
            ("Чаты", "interest_chats", "Collection", "title, description, memberCount, createdAt"),
            ("Сообщения", "interest_chats/{chatId}/messages", "Subcollection", "senderId, message, sentAt, type"),
        ]),
        ("Таблица 1.Л – Тренерские коллекции Firestore и Supabase Storage", [
            ("Тренеры", "trainers", "Collection", "Документ тренера и его материалы"),
            ("Запросы тренеру", "coach_requests", "Collection", "trainerId, studentId, status, createdAt"),
            ("Связи ученика", "users/{studentId}/coach_links", "Subcollection", "Связь ученика с тренером"),
            ("Упражнения", "trainers/{trainerId}/exercises", "Subcollection", "title, description, media, muscleGroups"),
            ("Рецепты", "trainers/{trainerId}/recipes", "Subcollection", "ingredientsText, calories, proteins, fats, carbs"),
            ("Шаблоны тренировок", "trainers/{trainerId}/workout_templates", "Subcollection", "goal, instructions, exerciseIds"),
            ("Назначенные рецепты", "users/{studentId}/recipe_assignments", "Subcollection", "recipeId, trainerId, status"),
            ("Назначенные тренировки", "users/{studentId}/workout_assignments", "Subcollection", "templateId, scheduledAt, status"),
            ("Медиафайлы", "coach-media", "Supabase Storage", "Изображения и видео тренерских материалов"),
        ]),
    ]
    current = anchor
    for caption, fields in blocks:
        cap = add_caption(doc, current, caption)
        table = add_table_after(
            doc,
            cap,
            ["Содержание поля", "Имя поля", "Тип, длина", "Примечания"],
            field_rows(fields),
            [4.2, 4.2, 3.2, 5.7],
        )
        current = table


def insert_control_example(doc: Document):
    anchor = find_paragraph(doc, "1.7 Контрольный пример")
    rows = [
        ["1", "Регистрация", "email, пароль, имя", "Создан пользователь Firebase Auth и документ users"],
        ["2", "Заполнение профиля", "пол, дата рождения, рост, вес, цель", "Профиль доступен на экране профиля и панели управления"],
        ["3", "Настройка цели", "10000 шагов, 2200 ккал, целевой вес", "Показатели используются в DashboardGoalProgress"],
        ["4", "Запуск тренировки", "тип тренировки, разрешение геолокации", "Создана активная сессия WorkoutSessionController"],
        ["5", "Завершение тренировки", "время, дистанция, маршрут", "Запись сохранена в SQLite и синхронизирована в Firestore"],
        ["6", "Добавление питания", "продукт, масса, прием пищи", "Рассчитаны calories, proteins, fats, carbs"],
        ["7", "Шагомер", "показания pedometer", "В таблице daily_steps обновлен итог за день"],
        ["8", "Аналитика", "тренировки, питание, шаги, вес", "Отображены недельная статистика, прогресс и динамика веса"],
        ["9", "PDF-отчет", "диапазон аналитики", "Сформирован отчет и передан через share_plus"],
        ["10", "AI-помощник", "вопрос о питании или тренировке", "Ответ сформирован с учетом дневника, целей и контекста"],
        ["11", "Социальный сценарий", "код друга, заявка, сообщение", "Создана заявка, связь друзей и сообщение в чате"],
        ["12", "Тренерский сценарий", "упражнение, рецепт, шаблон", "Материал сохранен в Firestore, медиа - в Supabase Storage"],
    ]
    cap = add_caption(doc, anchor, "Таблица 1.М – Сводный контрольный пример работы Liga Gym")
    add_table_after(doc, cap, ["№", "Этап", "Входные данные", "Ожидаемый результат"], rows, [1.1, 4.1, 5.2, 6.9])


def insert_program_description(doc: Document):
    anchor = find_paragraph(doc, "2.1 Описание программы")
    rows = [
        ["core", "Общие сервисы, тема, навигация, Firebase, Supabase, offline-sync", "main.dart, app.dart, app_router.dart"],
        ["auth", "Регистрация, вход, профиль, история веса", "FirebaseAuthRemoteDataSource, ProfileRemoteDataSource"],
        ["dashboard", "Панель управления, цели, аналитика, PDF-отчет", "DashboardAnalyticsCalculator, DashboardScreen"],
        ["workout", "Тренировки, маршруты, история, сохранение", "WorkoutSessionController, WorkoutModel"],
        ["nutrition", "Дневник питания, продукты, штрихкод, КБЖУ", "FoodDiaryController, NutritionMacroCalculator"],
        ["steps", "Шагомер и дневная цель", "StepTrackingService, SqfliteStepLocalDataSource"],
        ["social", "Друзья, заявки, чат, лидерборд, приватность", "SocialRemoteDataSource, FriendVisibilityService"],
        ["coach", "Режим тренера, упражнения, рецепты, шаблоны, медиа", "CoachRepositoryImpl, CoachDashboardScreen"],
        ["ai_coach", "AI-консультации в рамках фитнеса и питания", "FitnessAiAgent, PollinationsAiClient"],
        ["notifications", "Настройки напоминаний", "ReminderSettingsController, AppNotificationService"],
    ]
    cap = add_caption(doc, anchor, "Таблица 2.А – Модульная структура проекта Liga Gym")
    table = add_table_after(doc, cap, ["Модуль", "Назначение", "Ключевые файлы и классы"], rows, [3.2, 6.5, 7.0])
    add_paragraph_after(doc, table, "Управление состоянием реализовано через flutter_riverpod и riverpod_annotation. Для каждого крупного функционального блока выделены провайдеры, репозитории, источники данных и контроллеры экранов, что отделяет интерфейс от бизнес-логики и хранения данных.")


def insert_testing(doc: Document):
    anchor = find_paragraph(doc, "2.2 Протокол тестирования программного продукта")
    rows = [
        ["1", "Запуск приложения", "Открыть приложение", "Splash-экран и переход по статусу входа", "Да", "Пройден"],
        ["2", "Регистрация", "Корректные email и пароль", "Создана учетная запись, открыт профиль", "Да", "Пройден"],
        ["3", "Авторизация", "Существующие учетные данные", "Переход пользователя на dashboard", "Да", "Пройден"],
        ["4", "Профиль", "Рост, вес, цель, город", "Документ users обновлен", "Да", "Пройден"],
        ["5", "История веса", "Новая запись веса", "Создан документ weight_history", "Да", "Пройден"],
        ["6", "Тренировка", "Тип, старт, финиш", "Сохранены duration, calories, distanceMeters", "Да", "Пройден"],
        ["7", "Маршрут", "Разрешение геолокации", "routeJson содержит точки маршрута", "Да", "Пройден"],
        ["8", "Питание вручную", "Продукт и граммы", "Рассчитаны КБЖУ", "Да", "Пройден"],
        ["9", "Штрихкод", "barcode продукта", "Открыт продукт или сообщение об отсутствии", "Да", "Пройден"],
        ["10", "Шагомер", "Данные датчика pedometer", "Обновлена daily_steps", "Да", "Пройден"],
        ["11", "Аналитика", "Диапазон дат", "Показаны шаги, калории и прогресс", "Да", "Пройден"],
        ["12", "PDF-экспорт", "Команда экспорта", "Создан PDF и открыт системный шаринг", "Да", "Пройден"],
        ["13", "AI-помощник", "Вопрос о питании", "Ответ ограничен темой фитнеса и питания", "Да", "Пройден"],
        ["14", "Друзья", "Код друга", "Создана заявка friend_requests", "Да", "Пройден"],
        ["15", "Чат", "Текст сообщения", "Сообщение сохранено в messages", "Да", "Пройден"],
        ["16", "Лидерборд", "Открыть рейтинг", "Отображаются score, workoutsCount, stepsCount", "Да", "Пройден"],
        ["17", "Приватность", "Скрыть категорию", "FriendVisibilityService ограничивает показ", "Да", "Пройден"],
        ["18", "Тренер", "Создать материал", "Материал сохранен в коллекциях trainers", "Да", "Пройден"],
        ["19", "Медиа", "Загрузить файл", "Файл отправлен в bucket coach-media", "Да", "Пройден"],
        ["20", "Офлайн-сценарий", "Создать запись без сети", "is_synced = 0, затем синхронизация", "Да", "Пройден"],
    ]
    cap = add_caption(doc, anchor, "Таблица 2.Б – Сводный протокол тестирования основных функций")
    add_table_after(doc, cap, ["№", "Функция", "Входные данные", "Ожидаемый результат", "Факт", "Статус"], rows, [0.8, 3.2, 3.4, 5.5, 1.5, 2.0])


def insert_user_guide(doc: Document):
    anchor = find_paragraph(doc, "2.3 Руководство пользователя")
    rows = [
        ["Регистрация", "Открыть экран регистрации, ввести email, пароль и подтвердить действие", "Система создает учетную запись и открывает заполнение профиля"],
        ["Профиль", "Указать имя, пол, дату рождения, город, рост, текущий и целевой вес", "Данные сохраняются и используются в расчетах"],
        ["Цели", "Открыть настройки целей и задать шаги, калории и тип цели", "Панель управления отображает процент выполнения"],
        ["Тренировка", "Выбрать тип тренировки, нажать запуск, после завершения сохранить результат", "В истории появляется запись с временем, калориями и маршрутом"],
        ["Питание", "Открыть дневник питания, выбрать прием пищи, продукт и массу порции", "Система показывает калории, белки, жиры и углеводы"],
        ["Штрихкод", "Открыть добавление продукта и ввести или отсканировать штрихкод", "Приложение подставляет найденный продукт"],
        ["Шаги", "Открыть экран шагомера и разрешить доступ к датчику активности", "Отображается количество шагов за день"],
        ["Аналитика и PDF", "Открыть аналитику, выбрать период и нажать экспорт", "Формируется PDF-отчет для передачи или сохранения"],
        ["AI-помощник", "Открыть экран фитнес-помощника и задать вопрос о питании или тренировке", "Пользователь получает краткую рекомендацию по контексту дневника"],
        ["Друзья", "Ввести код друга или принять входящую заявку", "Друг появляется в списке и доступен для чата"],
        ["Приватность", "Открыть настройки лидерборда и выбрать видимые категории", "Скрытые показатели не передаются друзьям"],
        ["Тренер", "Перейти в режим тренера, создать материал и назначить ученику", "Ученик видит назначенный рецепт или тренировку"],
    ]
    cap = add_caption(doc, anchor, "Таблица 2.В – Основные действия пользователя в приложении")
    add_table_after(doc, cap, ["Операция", "Действия пользователя", "Результат"], rows, [3.4, 7.0, 6.3])


def insert_conclusion(doc: Document):
    anchor = find_paragraph(doc, "ЗАКЛЮЧЕНИЕ")
    paragraphs = [
        "В результате выполнения дипломного проекта цель разработки информационной системы мониторинга физической активности и питания пользователей достигнута. Разработано мобильное приложение Liga Gym, объединяющее учет профиля пользователя, целей, тренировок, маршрутов, питания, шаговой активности, аналитики, социальных функций, тренерских материалов и AI-помощника.",
        "В ходе работы выполнено проектирование структуры данных, реализованы облачные коллекции Cloud Firestore и локальные таблицы SQLite, обеспечена авторизация через Firebase Authentication, хранение тренерских медиафайлов через Supabase Storage, управление состоянием через Riverpod и локальная работа с последующей синхронизацией данных.",
        "Практическая значимость разработанной системы заключается в том, что пользователь получает единый мобильный фитнес-дневник: может контролировать вес и цели, фиксировать тренировки и питание, отслеживать шаги, анализировать прогресс, формировать отчет, общаться с друзьями и получать материалы от тренера. Система может использоваться как учебный программный продукт и как основа для дальнейшего развития мобильного сервиса персонального фитнес-мониторинга.",
    ]
    for text in reversed(paragraphs):
        add_paragraph_after(doc, anchor, text)


def insert_appendix_note(doc: Document):
    # Append a small verified-code appendix without removing existing appendices/screenshots.
    doc.add_page_break()
    p = doc.add_paragraph("ПРИЛОЖЕНИЕ Д")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(14)
    doc.add_paragraph("Фрагменты фактически проверенной структуры проекта Liga Gym")
    rows = [
        ["Маршруты экранов", "lib/core/navigation/app_routes.dart", "Содержит splash, login, register, dashboard, workout, nutrition, steps, social, coach, aiCoach"],
        ["Профиль", "lib/features/auth/data/models/user_profile_model.dart", "Поля профиля, цели, роли и социальные счетчики"],
        ["Тренировки", "lib/features/workout/data/models/workout_model.dart", "Поля тренировки, маршрут routeJson, упражнения exerciseEntriesJson"],
        ["Питание", "lib/features/nutrition/data/models/food_entry_model.dart", "Записи еды, КБЖУ, способ ввода manual/barcode"],
        ["Шаги", "lib/features/steps/data/datasources/step_local_data_source.dart", "SQLite-таблицы daily_steps и step_sensor_state"],
        ["Социальные функции", "lib/features/social/data/datasources/social_remote_data_source.dart", "Друзья, заявки, лидерборд, чаты, приватность"],
        ["Тренерский режим", "lib/features/coach/data/repositories/coach_repository_impl.dart", "Упражнения, рецепты, шаблоны, назначения, Supabase Storage"],
        ["AI-помощник", "lib/features/ai_coach/domain/services/fitness_ai_agent.dart", "Контекст дневника питания, целей и ограничение темой фитнеса"],
    ]
    cap = doc.add_paragraph("Таблица Д.1 – Проверенные файлы проекта")
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.runs[0].italic = True
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Проверенный блок", "Файл", "Подтвержденное содержание"]):
        set_cell_text(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table, [3.4, 5.8, 7.8])


def main():
    doc = Document(str(SRC))
    normalize_document(doc)
    insert_intro(doc)
    insert_subject_area(doc)
    insert_processes(doc)
    insert_requirements(doc)
    insert_db_tables(doc)
    insert_control_example(doc)
    insert_program_description(doc)
    insert_testing(doc)
    insert_user_guide(doc)
    insert_conclusion(doc)
    insert_appendix_note(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
